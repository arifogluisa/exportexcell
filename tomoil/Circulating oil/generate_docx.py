import os

import pandas as pd
from docx import Document
from docx.shared import Pt


def extract_data_from_excell(file_name: str, sheet_name: str) -> None:
    result = []
    df = pd.read_excel(file_name, sheet_name=sheet_name)
    for main_row_index in range(3, len(df), 9):
        main_index, main_row = tuple(df.iterrows())[main_row_index]
        pds_no = tuple(main_row.items())[1][1]
        product_name = tuple(main_row.items())[2][1]
        product_description = tuple(main_row.items())[3][1]
        benefits = tuple(main_row.items())[5][1]
        benefits = benefits.split("\n")
        oem_claims = tuple(main_row.items())[7][1]

        performance_claims = get_performance_claims(oem_claims)
        row_item = {
            "pds_no": pds_no,
            "product_name": product_name,
            "product_description": product_description,
            "benefits": benefits,
            "performance_claims": performance_claims,
        }
        result.append(row_item)
    return result


def get_performance_claims(
    oem_claims: str,
) -> list:
    result = []
    if oem_claims not in ["Not available", "Not applicable", "", None]:
        if "," in oem_claims:
            oem_claims = oem_claims.split(",")
        else:
            oem_claims = oem_claims.split("\n")
        result.extend([item.strip() for item in oem_claims])
    return result


def generate_docx(
    template_name: str,
    pds_no: str,
    product_name: str,
    product_description: str,
    fetaures: list,
    oil_type: str,
    performance_claims: list,
) -> None:
    doc = Document(template_name)
    product_name_place = "Product-name"
    product_description_place = "Prod-desc"
    fetaures_place = "Features-benefits"
    performance_claims_place = "Performance-claim"
    performance_claims_title = "Specifications"
    oil_type_place = "Oil-type"

    for paragraph in doc.paragraphs:
        if fetaures_place in paragraph.text:
            paragraph.clear()
            paragraph.text = "\n".join([f"• {item}" for item in fetaures])
            for run in paragraph.runs:
                run.font.name = "Arial"

        if performance_claims_place in paragraph.text:
            paragraph.clear()
            if performance_claims:
                paragraph.text = "\n".join([f"• {item}" for item in performance_claims])
                for run in paragraph.runs:
                    run.font.name = "Arial"
            else:
                for paragraph_inner in doc.paragraphs:
                    if performance_claims_title in paragraph_inner.text:
                        paragraph_inner.clear()

        if product_name_place in paragraph.text:
            paragraph.text = paragraph.text.replace(product_name_place, product_name)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(16)
                run.bold = True

        if product_description_place in paragraph.text:
            paragraph.text = paragraph.text.replace(
                product_description_place, product_description
            )
            for run in paragraph.runs:
                run.font.name = "Arial"

        if oil_type_place in paragraph.text:
            paragraph.text = paragraph.text.replace(oil_type_place, oil_type)
            for run in paragraph.runs:
                run.font.name = "Arial"

    save_folder = r"../docs/Circulating oil"
    product_folder = product_name.replace("/", "-")
    for i in range(1, 4):
        for section in doc.sections:
            header = section.header
            header.paragraphs[4].text = f"PDS No.: {pds_no}/0{i}"
            for run in header.paragraphs[4].runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)

        if not os.path.exists(f"{save_folder}/{product_folder}"):
            os.makedirs(f"{save_folder}/{product_folder}")

        doc.save(f"{save_folder}/{product_folder}/{product_folder}_0{i}_eng.docx")


excell_files = extract_data_from_excell("../exc_temp.xlsx", "Circulating oil")


for _ in excell_files:
    generate_docx(
        "tom_test.docx",
        _["pds_no"],
        _["product_name"],
        _["product_description"],
        _["benefits"],
        "Circulating oil",
        _["performance_claims"],
    )
