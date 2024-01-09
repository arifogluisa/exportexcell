import os

import pandas as pd
from docx import Document
from docx.shared import Pt

# def extract_data_from_excell(file_name: str, sheet_name: str) -> None:
#     result = []
#     df = pd.read_excel(file_name, sheet_name=sheet_name)
#     for main_row_index in range(3, len(df), 9):
#         main_index, main_row = tuple(df.iterrows())[main_row_index]
#         pds_no = tuple(main_row.items())[1][1]
#         product_name = tuple(main_row.items())[2][1]
#         product_description = tuple(main_row.items())[3][1]
#         benefits = tuple(main_row.items())[5][1]
#         benefits = benefits.split('\n')
#         sae_grade = tuple(main_row.items())[7][1]
#         api = tuple(main_row.items())[8][1]
#         ilsac = tuple(main_row.items())[9][1]
#         acea = tuple(main_row.items())[10][1]
#         oem_claims = tuple(main_row.items())[11][1]
#         recommendations = tuple(main_row.items())[12][1]

#         performance_claims = get_performance_claims(api, ilsac, acea, oem_claims, recommendations)
#         row_item = {
#             'pds_no': pds_no,
#             'product_name': product_name,
#             'product_description': product_description,
#             'benefits': benefits,
#             'sae_grade': sae_grade,
#             'performance_claims': performance_claims,
#         }
#         result.append(row_item)
#     return result


# def get_performance_claims(
#         api: str,
#         ilsac: str,
#         acea: str,
#         oem_claims: str,
#         recommendations: str
#         ) -> list:
#     result = []
#     if api not in ['Not available', 'Not applicable', '', None]:
#         result.append(f'API {api}')
#     if ilsac not in ['Not available', 'Not applicable', '', None]:
#         result.append(f'ILSAC {ilsac}')
#     if acea not in ['Not available', 'Not applicable', '', None]:
#         result.append(f'ACEA {acea}')
#     if oem_claims not in ['Not available', 'Not applicable', '', None]:
#         if ',' in oem_claims:
#             oem_claims = oem_claims.split(',')
#         else:
#             oem_claims = oem_claims.split('\n')
#         result.extend([item.strip() for item in oem_claims])
#     if recommendations not in ['Not available', 'Not applicable', '', None]:
#         if ',' in recommendations:
#             recommendations = recommendations.split(',')
#         else:
#             recommendations = recommendations.split('\n')
#         result.extend([item.strip() for item in recommendations])
#     return result


# def generate_docx(
#     template_name: str,
#     pds_no: str,
#     product_name: str,
#     product_description: str,
#     fetaures: list,
#     sae_grade: str,
#     oil_type: str,
#     performance_claims: list,
# ) -> None:
#     doc = Document(template_name)
#     product_name_place = "Product-name"
#     product_description_place = "Prod-desc"
#     fetaures_place = "Features-benefits"
#     performance_claims_place = "Performance-claim"
#     performance_claims_title = "Performance claims"
#     oil_type_place = "Oil-type"

#     for paragraph in doc.paragraphs:
#         if fetaures_place in paragraph.text:
#             paragraph.clear()
#             paragraph.text = "\n".join([f"• {item}" for item in fetaures])
#             for run in paragraph.runs:
#                 run.font.name = "Arial"

#         if performance_claims_place in paragraph.text:
#             paragraph.clear()
#             if performance_claims:
#                 paragraph.text = "\n".join([f"• {item}" for item in performance_claims])
#                 for run in paragraph.runs:
#                     run.font.name = "Arial"
#             else:
#                 for paragraph_inner in doc.paragraphs:
#                     if performance_claims_title in paragraph_inner.text:
#                         paragraph_inner.clear()

#         if product_name_place in paragraph.text:
#             paragraph.text = paragraph.text.replace(
#                 product_name_place, product_name
#             )
#             for run in paragraph.runs:
#                 run.font.name = "Arial"
#                 run.font.size = Pt(16)
#                 run.bold = True

#         if product_description_place in paragraph.text:
#             paragraph.text = paragraph.text.replace(
#                 product_description_place, product_description
#             )
#             for run in paragraph.runs:
#                 run.font.name = "Arial"

#         if oil_type_place in paragraph.text:
#             paragraph.text = paragraph.text.replace(oil_type_place, oil_type)
#             for run in paragraph.runs:
#                 run.font.name = "Arial"

#     save_folder = r"engine"
#     product_folder = product_name.replace("/", "-")
#     for i in range(1, 4):
#         for section in doc.sections:
#             header = section.header
#             header.paragraphs[4].text = f"PDS No.: {pds_no}/0{i}"
#             for run in header.paragraphs[4].runs:
#                 run.font.name = "Arial"
#                 run.font.size = Pt(9)
            
#         if not os.path.exists(f'{save_folder}/{product_folder}'):
#             os.makedirs(f'{save_folder}/{product_folder}')

#         doc.save(
#             f'{save_folder}/{product_folder}/{product_folder}_{sae_grade}_0{i}_eng.docx'
#         )


def extract_data_from_excell(file_name: str, sheet_name: str) -> None:
    result = []
    df = pd.read_excel(file_name, sheet_name=sheet_name)
    for main_row_index in range(3, len(df), 9):
        main_index, main_row = tuple(df.iterrows())[main_row_index]
        pds_no = tuple(main_row.items())[1][1]
        product_name = tuple(main_row.items())[2][1]
        product_description = tuple(main_row.items())[3][1]
        benefits = tuple(main_row.items())[5][1]
        benefits = benefits.split('\n')
        # sae_grade_value = tuple(main_row.items())[6][1]
        # sae_grade = sae_grade_value if sae_grade_value != 'Not available' else ''
        api = tuple(main_row.items())[7][1]
        # oem_claims = tuple(main_row.items())[6][1]
        specifications = tuple(main_row.items())[8][1]

        print(main_index, pds_no)

        performance_claims = get_performance_claims(api, specifications)
        row_item = {
            'pds_no': pds_no,
            'product_name': product_name,
            'product_description': product_description,
            'benefits': benefits,
            # 'sae_grade': sae_grade,
            'performance_claims': performance_claims,
        }
        result.append(row_item)
    return result


def get_performance_claims(
        # oem_claims: str,
        api: str,
        specifications: str
        ) -> list:
    result = []
    if api not in ['Not available', 'Not applicable', '', None]:
        result.append(f'API {api}')
    # if oem_claims not in ['Not available', 'Not applicable', '', None]:
    #     if ',' in oem_claims:
    #         oem_claims = oem_claims.split(',')
    #     else:
    #         oem_claims = oem_claims.split('\n')
    #     result.extend([item.strip() for item in oem_claims if item])
    if specifications not in ['Not available', 'Not applicable', '', None]:
        if ',' in specifications:
            specifications = specifications.split(',')
        else:
            specifications = specifications.split('\n')
        result.extend([item.strip() for item in specifications if item])
    return result


def generate_docx(
    template_name: str,
    pds_no: str,
    product_name: str,
    product_description: str,
    fetaures: list,
    oil_type: str,
    performance_claims: list,
) -> None:
    doc = Document(template_name)
    product_name_place = "Product-name"
    product_description_place = "Prod-desc"
    fetaures_place = "Features-benefits"
    performance_claims_place = "Performance-claim"
    performance_claims_title = "Performance claims"
    oil_type_place = "Oil-type"

    for paragraph in doc.paragraphs:
        if fetaures_place in paragraph.text:
            paragraph.clear()
            paragraph.text = "\n".join([f"• {item}" for item in fetaures])
            for run in paragraph.runs:
                run.font.name = "Arial"

        if performance_claims_place in paragraph.text:
            paragraph.clear()
            if performance_claims:
                paragraph.text = "\n".join([f"• {item}" for item in performance_claims])
                for run in paragraph.runs:
                    run.font.name = "Arial"
            else:
                for paragraph_inner in doc.paragraphs:
                    if performance_claims_title in paragraph_inner.text:
                        paragraph_inner.clear()

        if product_name_place in paragraph.text:
            paragraph.text = paragraph.text.replace(
                product_name_place, product_name
            )
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(16)
                run.bold = True

        if product_description_place in paragraph.text:
            paragraph.text = paragraph.text.replace(
                product_description_place, product_description
            )
            for run in paragraph.runs:
                run.font.name = "Arial"

        if oil_type_place in paragraph.text:
            paragraph.text = paragraph.text.replace(oil_type_place, oil_type)
            for run in paragraph.runs:
                run.font.name = "Arial"

    save_folder = r"Transmission"
    product_folder = product_name.replace("/", "-").strip()
    for i in range(1, 4):
        for section in doc.sections:
            header = section.header
            header.paragraphs[4].text = f"PDS No.: {pds_no}/0{i}"
            for run in header.paragraphs[4].runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)
            
        if not os.path.exists(f'{save_folder}/{product_folder}'):
            os.makedirs(f'{save_folder}/{product_folder}')

        doc.save(
            f'{save_folder}/{product_folder}/{product_folder}_0{i}_eng.docx'
        )


excell_files = extract_data_from_excell('excell_temp.xlsx', 'Transmission oil')


for _ in excell_files:
    generate_docx(
        'atf_temp.docx',
        _['pds_no'],
        _['product_name'],
        _['product_description'],
        _['benefits'],
        'Transmission oil',
        _['performance_claims'],
        )
