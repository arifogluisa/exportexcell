import os
from docx import Document
from docx.shared import Pt


def generate_docx(
    template_name: str,
    pds_no: str,
    product_name: str,
    product_description: str,
    fetaures: list,
    sae_grade: str,
    oil_type: str,
    performance_claims: list,
) -> None:
    doc = Document(template_name)
    pds_no_place = "pds_no"
    product_name_place = "Product-name"
    product_description_place = "Prod-desc"
    fetaures_place = "Features-benefits"
    performance_claims_place = "Performance-claims"
    performance_claims_title = "Performance claims"
    oil_type_place = "Oil-type"

    for paragraph in doc.paragraphs:
        if fetaures_place in paragraph.text:
            paragraph.clear()
            paragraph.text = "\n".join([f"• {item}" for item in fetaures])
            for run in paragraph.runs:
                run.font.name = "Arial"

        if performance_claims_place in paragraph.text:
            paragraph.clear()
            if performance_claims:
                paragraph.text = "\n".join([f"• {item}" for item in performance_claims])
                for run in paragraph.runs:
                    run.font.name = "Arial"
            else:
                for paragraph_inner in doc.paragraphs:
                    if performance_claims_title in paragraph_inner.text:
                        paragraph_inner.clear()

        if product_name_place in paragraph.text:
            paragraph.text = paragraph.text.replace(
                product_name_place, f"{product_name} {sae_grade}"
            )
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(16)
                run.bold = True

        if product_description_place in paragraph.text:
            paragraph.text = paragraph.text.replace(
                product_description_place, product_description
            )
            for run in paragraph.runs:
                run.font.name = "Arial"

        if oil_type_place in paragraph.text:
            paragraph.text = paragraph.text.replace(oil_type_place, oil_type)
            for run in paragraph.runs:
                run.font.name = "Arial"

    save_folder = r"engine_oils"
    product_folder = product_name.replace("/", "-")
    for i in range(1, 4):
        for section in doc.sections:
            header = section.header
            header.paragraphs[4].text = f"PDS No.: {pds_no}/0{i}"
            for run in header.paragraphs[4].runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)
            
        if not os.path.exists(f'{save_folder}/{product_folder}'):
            os.makedirs(f'{save_folder}/{product_folder}')

        doc.save(
            f'{save_folder}/{product_folder}/{product_folder}_{sae_grade}_0{i}_eng.docx'
        )
